"""
We detect file type from MAGIC BYTES (first 8 bytes of the file),
    NOT from the filename extension. Why? An attacker could name a
    malicious file "resume.pdf" when it's actually an executable.
    Magic bytes can't be faked without breaking the file format.

MAGIC BYTES:
    PDF:  %PDF  (hex: 25 50 44 46)
    DOCX: PK    (hex: 50 4B) — it's a ZIP archive
    The 'content_type' from the upload is also checked but not trusted alone.
    """



from __future__ import annotations


import io
import re
import unicodedata
from pathlib import Path

import fitz #PyMUPdf
import structlog
from docx import Document

logger=structlog.get_logger()
# Maximum characters we'll return. Prevents abuse and fits in
# embedding context windows (text-embedding-3-small handles 8191 tokens).
MAX_TEXT_LENGTH=8000

PDF_MAGIC=b"%PDF"
DOCX_MAGIC=b"PK" #docx is a zip archive starts with PK

def detect_file_type(content: bytes)-> str:
    """
    detects file type from magic bytes(first few bytes of file)
    return pdf docx or unknown
    """

    if content[:4] == PDF_MAGIC:
        return "pdf"
    if content[:2] == DOCX_MAGIC:
        return "docx"
    return "unknown"

def clean_text(text:str)-> str:
    """
    Normalize and clean extracted text.
    Steps:
    1. Unicode normalization (NFC) — handles accented chars consistently
    2. Replace multiple whitespace/newlines with single space
    3. Strip leading/trailing whitespace
    4. Truncate to MAX_TEXT_LENGTH:
    """

    text=unicodedata.normalize("NFC",text)
    text=re.sub(r"\s+"," ",text)
    text=text.strip()
    if len(text) > MAX_TEXT_LENGTH:
        text=text[:MAX_TEXT_LENGTH]
    return text

def extract_from_pdf_bytes(content: bytes)-> str:
    # we raise value error if exracted text too small <50 chars likely
    # a scanned image PDF withour OCR
    doc=fitz.open(stream=content,filetype="pdf")
    pages_text=[]

    for page_num in range(len(doc)):
        page=doc[page_num]
        text=page.get_text("text") # "text" mode = reading order
        if text.strip():
            pages_text.append(text)

    doc.close()

    full_text="\n\n".join(pages_text)
    cleaned=clean_text(full_text)

    if len(cleaned) < 50:
        raise ValueError(
            f"Extracted only {len(cleaned)} chars from PDF. "
            "The file may be a scanned image (OCR not supported)."
        )
    logger.debug("pdf_extracted", pages=len(pages_text), chars=len(cleaned))
    return cleaned


def extract_from_docx_bytes(content: bytes) -> str:
    """
    Extract text from a DOCX file's raw bytes.
    DOCX files are ZIP archives containing XML. python-docx parses
    the XML tree. We extract from two sources:
    1. Paragraphs — normal text content
    2. Tables — structured data (skills tables, experience tables)
    Args:
        content: Raw bytes of the DOCX file
    Returns:
        Cleaned text string
    """
    doc = Document(io.BytesIO(content))
    parts = []
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Extract table content (resumes often have skills in tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    full_text = "\n".join(parts)
    cleaned = clean_text(full_text)
    logger.debug("docx_extracted", paragraphs=len(parts), chars=len(cleaned))
    return cleaned
def extract_text_from_bytes(content: bytes, filename: str = "") -> str:
    """
    Main entry point — detect file type and extract text.
    Args:
        content: Raw file bytes
        filename: Original filename (for logging only, NOT for type detection)
    Returns:
        Cleaned text string (max 8000 chars)
    Raises:
        ValueError: If file type is unsupported or text too short
    """
    file_type = detect_file_type(content)
    if file_type == "pdf":
        text = extract_from_pdf_bytes(content)
    elif file_type == "docx":
        text = extract_from_docx_bytes(content)
    elif content[:3] == b"\xef\xbb\xbf" or content[:100].decode("utf-8", errors="ignore").isprintable():
        # UTF-8 BOM or printable text — treat as plain text
        text = clean_text(content.decode("utf-8", errors="replace"))
    else:
        raise ValueError(
            f"Unsupported file type. Expected PDF or DOCX, got magic bytes: {content[:4].hex()}"
        )
    logger.info(
        "text_extracted",
        filename=filename,
        file_type=file_type,
        chars=len(text),
    )
    return text
def extract_text_from_file(file_path: str | Path) -> str:
    """
    Convenience function — read a file from disk and extract text.
    Used by the seed script to process resume files in bulk.
    Args:
        file_path: Path to the file on disk
    Returns:
        Cleaned text string
    """
    path = Path(file_path)
    content = path.read_bytes()
    return extract_text_from_bytes(content, filename=path.name)








